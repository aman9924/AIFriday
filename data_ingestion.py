import pandas as pd
import re
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document
import pdfplumber

import config

class DataIngestion:
    """Load and harmonize data from multiple sources."""
    
    def __init__(self):
        self.kpi_df = None
        self.findings_df = None
        self.painpoints_df = None
        self.benchmarks = {}
    
    # ===== Crystallus KPI Loading =====
    def load_crystallus_kpis(self) -> pd.DataFrame:
        """Load KPIs from Crystallus Excel workbook."""
        print(f"Loading KPIs from {config.CRYSTALLUS_KPI_FILE}")
        
        try:
            xl = pd.ExcelFile(config.CRYSTALLUS_KPI_FILE)
            all_kpis = []
            
            for sheet_name in xl.sheet_names:
                if any(vs in sheet_name for vs in ["O2C", "P2P", "P2M", "R2R", "H2R", "A2D", "R2S", "I2M", "I2C", "DM"]):
                    df = pd.read_excel(xl, sheet_name=sheet_name)
                    
                    # Extract value stream code from sheet name
                    vs_code = None
                    for code in ["O2C", "P2P", "P2M", "R2R", "H2R", "A2D", "R2S", "I2M", "I2C", "DM"]:
                        if code in sheet_name:
                            vs_code = code
                            break
                    
                    if vs_code and len(df) > 0:
                        # Standardize column names
                        df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]
                        df["vs_code"] = vs_code
                        df["value_stream"] = sheet_name
                        all_kpis.append(df)
            
            if all_kpis:
                self.kpi_df = pd.concat(all_kpis, ignore_index=True)
                self._standardize_kpi_columns()
                print(f"  Loaded {len(self.kpi_df)} KPIs from {len(all_kpis)} sheets")
            else:
                self.kpi_df = self._create_sample_kpis()
                
        except FileNotFoundError:
            print(f"  Warning: File not found. Using sample KPIs.")
            self.kpi_df = self._create_sample_kpis()
        
        return self.kpi_df
    
    def _standardize_kpi_columns(self):
        """Standardize KPI dataframe columns."""
        column_mapping = {
            "kpi": "kpi_name",
            "kpi_name": "kpi_name",
            "metric": "kpi_name",
            "measure": "kpi_name",
            "description": "definition",
            "definition": "definition",
            "current": "current_value",
            "current_value": "current_value",
            "as-is": "current_value",
            "target": "target_value",
            "target_value": "target_value",
            "to-be": "target_value",
            "benchmark": "target_value",
            "formula": "formula",
            "calculation": "formula",
            "source": "data_source",
            "data_source": "data_source",
            "owner": "process_owner",
            "process_owner": "process_owner"
        }
        
        new_columns = {}
        for col in self.kpi_df.columns:
            col_lower = col.lower().replace(" ", "_")
            if col_lower in column_mapping:
                new_columns[col] = column_mapping[col_lower]
        
        self.kpi_df = self.kpi_df.rename(columns=new_columns)
        
        # Ensure required columns exist
        for req_col in ["kpi_name", "definition", "current_value", "target_value", "vs_code"]:
            if req_col not in self.kpi_df.columns:
                self.kpi_df[req_col] = ""
        
        # Extract numeric values
        self.kpi_df["current_numeric"] = self.kpi_df["current_value"].apply(self._extract_numeric)
        self.kpi_df["target_numeric"] = self.kpi_df["target_value"].apply(self._extract_numeric)
        
        # Determine direction (higher/lower is better)
        self.kpi_df["target_direction"] = self.kpi_df["kpi_name"].apply(self._infer_direction)
    
    def _extract_numeric(self, value) -> float:
        """Extract numeric value from string."""
        if pd.isna(value):
            return None
        try:
            if isinstance(value, (int, float)):
                return float(value)
            text = str(value)
            # Remove % and extract number
            text = text.replace("%", "").replace(",", "").replace("$", "")
            numbers = re.findall(r"[-+]?\\d*\\.?\\d+", text)
            if numbers:
                return float(numbers[0])
        except:
            pass
        return None
    
    def _infer_direction(self, kpi_name: str) -> str:
        """Infer whether higher or lower is better."""
        lower_better = ["cycle time", "days", "error", "defect", "variance", "cost", "lead time", "aging", "overdue"]
        kpi_lower = str(kpi_name).lower()
        for term in lower_better:
            if term in kpi_lower:
                return "lower"
        return "higher"
    
    def _create_sample_kpis(self) -> pd.DataFrame:
        """Create sample KPIs for demonstration."""
        sample_kpis = [
            {"vs_code": "O2C", "kpi_name": "Days Sales Outstanding (DSO)", "current_value": "52 days", "target_value": "45 days", "definition": "Average days to collect receivables", "target_direction": "lower"},
            {"vs_code": "O2C", "kpi_name": "Perfect Order Rate", "current_value": "87%", "target_value": "95%", "definition": "Orders delivered complete, on-time, damage-free", "target_direction": "higher"},
            {"vs_code": "O2C", "kpi_name": "Order-to-Cash Cycle Time", "current_value": "18 days", "target_value": "12 days", "definition": "Days from order receipt to payment", "target_direction": "lower"},
            {"vs_code": "P2P", "kpi_name": "PR to PO Cycle Time", "current_value": "5 days", "target_value": "2 days", "definition": "Days from requisition to purchase order", "target_direction": "lower"},
            {"vs_code": "P2P", "kpi_name": "3-Way Match Rate", "current_value": "78%", "target_value": "95%", "definition": "Auto-matched PO/GR/Invoice rate", "target_direction": "higher"},
            {"vs_code": "P2M", "kpi_name": "OEE (Overall Equipment Effectiveness)", "current_value": "42%", "target_value": "65%", "definition": "Availability x Performance x Quality", "target_direction": "higher"},
            {"vs_code": "P2M", "kpi_name": "Forecast Accuracy", "current_value": "68%", "target_value": "85%", "definition": "Demand forecast vs actual", "target_direction": "higher"},
            {"vs_code": "R2R", "kpi_name": "Days to Close", "current_value": "8 days", "target_value": "5 days", "definition": "Days to complete period close", "target_direction": "lower"},
            {"vs_code": "R2R", "kpi_name": "Manual Journal Rate", "current_value": "35%", "target_value": "15%", "definition": "Percentage of manual journal entries", "target_direction": "lower"},
            {"vs_code": "H2R", "kpi_name": "Time to Hire", "current_value": "45 days", "target_value": "30 days", "definition": "Days from requisition to start date", "target_direction": "lower"},
            {"vs_code": "DM", "kpi_name": "Master Data Quality Score", "current_value": "72%", "target_value": "95%", "definition": "Completeness and accuracy of master data", "target_direction": "higher"},
        ]
        df = pd.DataFrame(sample_kpis)
        df["current_numeric"] = df["current_value"].apply(self._extract_numeric)
        df["target_numeric"] = df["target_value"].apply(self._extract_numeric)
        return df
    
    # ===== SAP Health Check Loading =====
    def load_sap_healthcheck(self) -> pd.DataFrame:
        """Load findings from SAP health check report."""
        print(f"Loading SAP health check from {config.ERP_HEALTHCHECK_FILE}")
        
        findings = []
        
        try:
            doc = Document(config.ERP_HEALTHCHECK_FILE)
            current_section = "General"
            current_severity = "Medium"
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect section headers
                if para.style.name.startswith("Heading"):
                    current_section = text
                    continue
                
                # Detect severity keywords
                text_lower = text.lower()
                if "critical" in text_lower:
                    current_severity = "Critical"
                elif "high risk" in text_lower or "high priority" in text_lower:
                    current_severity = "High"
                elif "medium" in text_lower:
                    current_severity = "Medium"
                elif "low" in text_lower:
                    current_severity = "Low"
                
                # Extract findings
                if any(kw in text_lower for kw in ["finding", "issue", "risk", "recommendation", "observation", "gap"]):
                    finding = {
                        "finding_id": f"F{len(findings)+1:03d}",
                        "title": text[:100],
                        "detail": text,
                        "category": self._categorize_finding(text),
                        "severity": current_severity,
                        "section": current_section,
                        "vs_code": self._map_finding_to_vs(text),
                        "module": self._extract_module(text)
                    }
                    findings.append(finding)
            
            print(f"  Extracted {len(findings)} findings")
            
        except FileNotFoundError:
            print("  Warning: File not found. Using sample findings.")
            findings = self._create_sample_findings()
        
        self.findings_df = pd.DataFrame(findings)
        return self.findings_df
    
    def _categorize_finding(self, text: str) -> str:
        """Categorize finding based on keywords."""
        text_lower = text.lower()
        categories = {
            "Security": ["security", "authorization", "access", "password", "vulnerability", "sap*"],
            "Performance": ["performance", "slow", "timeout", "memory", "batch", "runtime"],
            "Data Quality": ["data quality", "duplicate", "incomplete", "master data", "consistency"],
            "Customization": ["custom", "z-program", "modification", "enhancement", "abap"],
            "Compliance": ["compliance", "audit", "sox", "gdpr", "regulation", "gxp"],
            "Integration": ["interface", "integration", "rfc", "idoc", "api", "middleware"],
            "Infrastructure": ["infrastructure", "server", "database", "storage", "backup"]
        }
        
        for category, keywords in categories.items():
            if any(kw in text_lower for kw in keywords):
                return category
        return "General"
    
    def _map_finding_to_vs(self, text: str) -> str:
        """Map finding to value stream."""
        text_lower = text.lower()
        mappings = {
            "O2C": ["sales", "billing", "receivable", "credit", "order", "delivery", "pricing"],
            "P2P": ["procurement", "purchase", "vendor", "invoice", "payment", "supplier"],
            "P2M": ["production", "manufacturing", "mrp", "plant", "quality", "batch"],
            "R2R": ["finance", "accounting", "ledger", "close", "consolidation", "journal"],
            "H2R": ["hr", "payroll", "employee", "personnel", "recruitment", "training"],
            "A2D": ["asset", "depreciation", "maintenance", "equipment", "capex"],
            "DM": ["master data", "material master", "customer master", "vendor master", "business partner"]
        }
        
        for vs_code, keywords in mappings.items():
            if any(kw in text_lower for kw in keywords):
                return vs_code
        return "CROSS"
    
    def _extract_module(self, text: str) -> str:
        """Extract SAP module from text."""
        modules = ["FI", "CO", "SD", "MM", "PP", "QM", "PM", "HR", "PS", "WM", "LE", "BASIS", "ABAP", "BW", "CRM", "SRM", "GRC"]
        text_upper = text.upper()
        for module in modules:
            if module in text_upper:
                return module
        return "GENERAL"
    
    def _create_sample_findings(self) -> List[Dict]:
        """Create sample findings."""
        return [
            {"finding_id": "F001", "title": "Critical security vulnerabilities in RFC gateway", "detail": "RFC gateway access not restricted", "category": "Security", "severity": "Critical", "vs_code": "CROSS", "module": "BASIS"},
            {"finding_id": "F002", "title": "High number of custom Z-programs", "detail": "Over 500 custom programs require S/4HANA compatibility review", "category": "Customization", "severity": "High", "vs_code": "CROSS", "module": "ABAP"},
            {"finding_id": "F003", "title": "Support packages behind current level", "detail": "System is 3 support packages behind recommended level", "category": "Infrastructure", "severity": "Medium", "vs_code": "CROSS", "module": "BASIS"},
            {"finding_id": "F004", "title": "Duplicate vendor master records", "detail": "15% duplicate rate in vendor master data", "category": "Data Quality", "severity": "High", "vs_code": "P2P", "module": "MM"},
            {"finding_id": "F005", "title": "Manual pricing overrides in SD", "detail": "High rate of manual pricing changes impacting margin analysis", "category": "Compliance", "severity": "Medium", "vs_code": "O2C", "module": "SD"},
        ]
    
    # ===== Pain Points Loading =====
    def load_painpoints(self) -> pd.DataFrame:
        """Load business pain points from documents."""
        print(f"Loading pain points...")
        
        painpoints = []
        
        try:
            doc = Document(config.ERP_HEALTHCHECK_FILE)
            for para in doc.paragraphs:
                text = para.text.strip()
                text_lower = text.lower()
                
                if any(kw in text_lower for kw in ["pain point", "challenge", "issue", "problem", "difficulty", "frustration"]):
                    painpoint = {
                        "pp_id": f"PP{len(painpoints)+1:03d}",
                        "description": text,
                        "vs_code": self._map_finding_to_vs(text),
                        "lob": self._extract_lob(text),
                        "severity": self._assess_painpoint_severity(text)
                    }
                    painpoints.append(painpoint)
                    
        except FileNotFoundError:
            painpoints = self._create_sample_painpoints()
        
        self.painpoints_df = pd.DataFrame(painpoints)
        print(f"  Loaded {len(self.painpoints_df)} pain points")
        return self.painpoints_df
    
    def _extract_lob(self, text: str) -> str:
        """Extract line of business from text."""
        lobs = {
            "Commercial": ["commercial", "sales", "marketing"],
            "Manufacturing": ["manufacturing", "production", "plant"],
            "Finance": ["finance", "accounting", "treasury"],
            "Supply Chain": ["supply chain", "logistics", "warehouse"],
            "IT": ["it", "technology", "system"],
            "HR": ["hr", "human resources", "employee"]
        }
        
        text_lower = text.lower()
        for lob, keywords in lobs.items():
            if any(kw in text_lower for kw in keywords):
                return lob
        return "General"
    
    def _assess_painpoint_severity(self, text: str) -> str:
        """Assess pain point severity."""
        text_lower = text.lower()
        if any(kw in text_lower for kw in ["critical", "urgent", "severe", "major blocker"]):
            return "Critical"
        elif any(kw in text_lower for kw in ["significant", "high impact", "important"]):
            return "High"
        elif any(kw in text_lower for kw in ["moderate", "medium"]):
            return "Medium"
        return "Low"
    
    def _create_sample_painpoints(self) -> List[Dict]:
        """Create sample pain points."""
        return [
            {"pp_id": "PP001", "description": "Manual batch release process takes 3-5 days", "vs_code": "P2M", "lob": "Manufacturing", "severity": "High"},
            {"pp_id": "PP002", "description": "No visibility into real-time inventory across plants", "vs_code": "P2M", "lob": "Supply Chain", "severity": "Critical"},
            {"pp_id": "PP003", "description": "Month-end close requires extensive manual reconciliation", "vs_code": "R2R", "lob": "Finance", "severity": "High"},
            {"pp_id": "PP004", "description": "Customer order status not visible to sales team", "vs_code": "O2C", "lob": "Commercial", "severity": "Medium"},
        ]
    
    # ===== Save Curated Data =====
    def save_curated_data(self):
        """Save all curated data to parquet files."""
        if self.kpi_df is not None:
            self.kpi_df.to_parquet(config.KPI_TARGETS_FILE, index=False)
            print(f"Saved KPIs to {config.KPI_TARGETS_FILE}")
        
        if self.findings_df is not None:
            self.findings_df.to_parquet(config.FINDINGS_FILE, index=False)
            print(f"Saved findings to {config.FINDINGS_FILE}")
        
        if self.painpoints_df is not None:
            self.painpoints_df.to_parquet(config.PAINPOINTS_FILE, index=False)
            print(f"Saved painpoints to {config.PAINPOINTS_FILE}")
    
    # ===== Main Ingestion Pipeline =====
    def run_ingestion(self):
        """Run complete data ingestion pipeline."""
        print("=" * 60)
        print("ZeroRisk ERP Health Check - Data Ingestion")
        print("=" * 60)
        
        print("\\n1. Loading Crystallus KPIs...")
        self.load_crystallus_kpis()
        
        print("\\n2. Loading SAP Health Check findings...")
        self.load_sap_healthcheck()
        
        print("\\n3. Loading pain points...")
        self.load_painpoints()
        
        print("\\n4. Saving curated data...")
        self.save_curated_data()
        
        print("\\n" + "=" * 60)
        print("Data Ingestion Complete!")
        print("=" * 60)
        
        return self


if __name__ == "__main__":
    ingestion = DataIngestion()
    ingestion.run_ingestion()